"""
api.py - PrivacyGate Background Service (FastAPI) - Optimized
"""

import os
import sys
import uuid
import time
import shutil
import sqlite3
import tempfile
import datetime
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

BASE_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE_DIR))

from scanner.regex_scanner import scan_text_regex
from scanner.ner_scanner   import scan_text_ner
from scanner.context_rules import apply_context_rules
from redactor.text_redactor  import redact_text, redact_docx, redact_pdf
from redactor.xlsx_redactor  import redact_xlsx, get_full_text_xlsx
from redactor.pptx_redactor  import redact_pptx, get_full_text_pptx
from redactor.image_redactor import redact_image, extract_text_image

# ── Database ──────────────────────────────────────────────

DB_DIR  = Path(os.environ.get("PROGRAMDATA", "C:/ProgramData")) / "PrivacyGate"
DB_PATH = DB_DIR / "scans.db"

def _get_db():
    DB_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    _init_db(conn)
    return conn

def _init_db(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS scans (
            id            TEXT PRIMARY KEY,
            timestamp     TEXT NOT NULL,
            source        TEXT NOT NULL,
            file_name     TEXT,
            file_type     TEXT,
            original_size INTEGER,
            findings_count INTEGER NOT NULL DEFAULT 0,
            types_found   TEXT,
            high_count    INTEGER DEFAULT 0,
            med_count     INTEGER DEFAULT 0,
            low_count     INTEGER DEFAULT 0,
            duration_ms   INTEGER,
            status        TEXT DEFAULT 'ok'
        )
    """)
    conn.commit()

def _log_scan(source, findings, file_name=None, file_type=None,
              original_size=None, duration_ms=None, status="ok"):
    import json
    scan_id   = str(uuid.uuid4())
    timestamp = datetime.datetime.now().isoformat(timespec="seconds")
    types_found = list(set(f.get("type", "") for f in findings))
    high = sum(1 for f in findings if f.get("risk") == "HIGH")
    med  = sum(1 for f in findings if f.get("risk") == "MED")
    low  = sum(1 for f in findings if f.get("risk") == "LOW")
    conn = _get_db()
    try:
        conn.execute("""
            INSERT INTO scans
              (id, timestamp, source, file_name, file_type, original_size,
               findings_count, types_found, high_count, med_count, low_count,
               duration_ms, status)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (scan_id, timestamp, source, file_name, file_type, original_size,
              len(findings), json.dumps(types_found), high, med, low, duration_ms, status))
        conn.commit()
    finally:
        conn.close()
    return scan_id

# ── Temp files ────────────────────────────────────────────

TEMP_DIR = Path(tempfile.gettempdir()) / "privacygate_cleaned"
TEMP_DIR.mkdir(exist_ok=True)

def _temp_output_path(original_name: str) -> Path:
    ext  = Path(original_name).suffix.lower()
    return TEMP_DIR / f"pg_{uuid.uuid4().hex}{ext}"

def _cleanup_old_temp_files():
    cutoff = time.time() - 600
    try:
        for f in TEMP_DIR.iterdir():
            if f.is_file() and f.stat().st_mtime < cutoff:
                f.unlink(missing_ok=True)
    except Exception:
        pass

# ── File type routing ─────────────────────────────────────

TEXT_EXTENSIONS = {
    ".txt", ".csv", ".json", ".jsonl", ".py", ".js", ".ts",
    ".html", ".htm", ".xml", ".md", ".yaml", ".yml",
    ".ini", ".cfg", ".env", ".sh", ".bat", ".sql",
    ".log", ".jsx", ".tsx", ".css", ".scss",
}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".gif"}

def _get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in (".docx", ".doc"): return "docx"
    if ext == ".pdf":            return "pdf"
    if ext in (".xlsx", ".xlsm"): return "xlsx"
    if ext == ".pptx":           return "pptx"
    if ext in IMAGE_EXTENSIONS:  return "image"
    if ext in TEXT_EXTENSIONS:   return "text"
    return "text"

def _extract_text_from_file(file_path: Path, file_type: str) -> str:
    if file_type == "text":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    if file_type == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except ImportError:
            raise HTTPException(500, "pip install pdfplumber")
    if file_type == "docx":
        try:
            import docx as docxlib
            doc = docxlib.Document(str(file_path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except ImportError:
            raise HTTPException(500, "pip install python-docx")
    if file_type == "xlsx": return get_full_text_xlsx(str(file_path))
    if file_type == "pptx": return get_full_text_pptx(str(file_path))
    # NOTE: images handled separately — no call here
    return ""

def _run_scanner(text: str, sensitivity: str = "standard") -> list:
    findings = []
    findings.extend(scan_text_regex(text))
    findings.extend(scan_text_ner(text))
    if sensitivity == "deep":
        findings = apply_context_rules(findings, text)
    seen, unique = set(), []
    for f in findings:
        key = f.get("value", "").lower()
        if key and key not in seen:
            seen.add(key)
            unique.append(f)
    return unique

def _redact_file(file_path: Path, file_type: str, findings: list,
                 out_path: Path, ocr_text: str = None) -> Path:
    """
    Route to correct redactor. For images, ocr_text is already extracted
    so we skip a second OCR pass inside redact_image.
    """
    if not findings:
        shutil.copy2(str(file_path), str(out_path))
        return out_path

    if file_type == "text":
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        cleaned_text, _, _ = redact_text(text, findings)
        out_path.write_text(cleaned_text, encoding="utf-8")
        return out_path
    if file_type == "pdf":
        redact_pdf(str(file_path), findings, str(out_path))
        return out_path
    if file_type == "docx":
        redact_docx(str(file_path), findings, str(out_path))
        return out_path
    if file_type == "xlsx":
        redact_xlsx(str(file_path), findings, str(out_path))
        return out_path
    if file_type == "pptx":
        redact_pptx(str(file_path), findings, str(out_path))
        return out_path
    if file_type == "image":
        # Pass findings directly — image_redactor uses them to black out regions
        # OCR already done above, no second pass needed
        redact_image(str(file_path), findings, str(out_path))
        return out_path

    shutil.copy2(str(file_path), str(out_path))
    return out_path

# ── FastAPI app ───────────────────────────────────────────

app = FastAPI(
    title="PrivacyGate",
    description="Local privacy scanning service.",
    version="1.0.0",
    docs_url=None,
    redoc_url=None,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TextScanRequest(BaseModel):
    text: str
    sensitivity: Optional[str] = "standard"

class TextScanResponse(BaseModel):
    scan_id:        str
    findings_count: int
    findings:       list
    cleaned_text:   str
    types_found:    list
    high_count:     int
    med_count:      int
    low_count:      int
    duration_ms:    int

@app.get("/status")
def status():
    return {"status": "running", "service": "PrivacyGate", "version": "1.0.0"}

@app.post("/scan-text", response_model=TextScanResponse)
def scan_text_endpoint(req: TextScanRequest):
    if not req.text or not req.text.strip():
        raise HTTPException(400, "Text is empty.")
    start    = time.time()
    findings = _run_scanner(req.text, req.sensitivity or "standard")
    cleaned_text, _, _ = redact_text(req.text, findings)
    duration_ms = int((time.time() - start) * 1000)
    scan_id = _log_scan(source="text", findings=findings, duration_ms=duration_ms)
    types_found = list(set(f.get("type", "") for f in findings))
    high = sum(1 for f in findings if f.get("risk") == "HIGH")
    med  = sum(1 for f in findings if f.get("risk") == "MED")
    low  = sum(1 for f in findings if f.get("risk") == "LOW")
    return TextScanResponse(
        scan_id=scan_id, findings_count=len(findings), findings=findings,
        cleaned_text=cleaned_text, types_found=types_found,
        high_count=high, med_count=med, low_count=low, duration_ms=duration_ms,
    )

@app.post("/scan-file")
async def scan_file_endpoint(
    file:        UploadFile = File(...),
    sensitivity: str        = Form("standard"),
):
    _cleanup_old_temp_files()

    filename  = file.filename or "upload"
    file_type = _get_file_type(filename)
    suffix    = Path(filename).suffix.lower()
    tmp_in    = Path(tempfile.mktemp(suffix=suffix, dir=str(TEMP_DIR)))
    out_path  = _temp_output_path(filename)

    try:
        content = await file.read()
        tmp_in.write_bytes(content)
        original_size = len(content)

        start = time.time()

        # ── IMAGES: single OCR pass ───────────────────────
        # Extract text once, reuse for both scanning and redaction.
        # Old code called extract_text_image() here AND inside redact_image()
        # = 2x OCR. Now we do it once and pass findings directly.
        if file_type == "image":
            try:
                text = extract_text_image(str(tmp_in))
            except Exception as e:
                raise HTTPException(422, f"OCR failed: {e}")
        else:
            try:
                text = _extract_text_from_file(tmp_in, file_type)
            except Exception as e:
                raise HTTPException(422, f"Could not read file: {e}")

        # Scan extracted text
        findings = _run_scanner(text, sensitivity)

        # Redact — for images, findings already have the PII values
        # redact_image uses them to locate and black out text regions
        try:
            _redact_file(tmp_in, file_type, findings, out_path)
        except Exception as e:
            raise HTTPException(500, f"Redaction failed: {e}")

        duration_ms = int((time.time() - start) * 1000)

        scan_id = _log_scan(
            source="file", findings=findings, file_name=filename,
            file_type=file_type, original_size=original_size, duration_ms=duration_ms,
        )

        types_found = list(set(f.get("type", "") for f in findings))
        high = sum(1 for f in findings if f.get("risk") == "HIGH")
        med  = sum(1 for f in findings if f.get("risk") == "MED")
        low  = sum(1 for f in findings if f.get("risk") == "LOW")

        findings_summary = [
            {"type": f.get("type"), "replace": f.get("replace"), "risk": f.get("risk")}
            for f in findings
        ]

        import json
        return FileResponse(
            path=str(out_path),
            filename=f"PrivacyGate_Cleaned_{filename}",
            media_type="application/octet-stream",
            headers={
                "X-Scan-Id":           scan_id,
                "X-Findings-Count":    str(len(findings)),
                "X-Types-Found":       ",".join(types_found),
                "X-High-Count":        str(high),
                "X-Med-Count":         str(med),
                "X-Low-Count":         str(low),
                "X-Duration-Ms":       str(duration_ms),
                "X-Original-Filename": filename,
                "X-Findings-Summary":  json.dumps(findings_summary),
                "Access-Control-Expose-Headers": (
                    "X-Scan-Id, X-Findings-Count, X-Types-Found, "
                    "X-High-Count, X-Med-Count, X-Low-Count, "
                    "X-Duration-Ms, X-Original-Filename, X-Findings-Summary"
                ),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Unexpected error: {e}")
    finally:
        try:
            tmp_in.unlink(missing_ok=True)
        except Exception:
            pass

@app.get("/stats")
def get_stats():
    import json
    conn = _get_db()
    try:
        row = conn.execute("""
            SELECT COUNT(*) AS total_scans, SUM(findings_count) AS total_pii_blocked,
                   SUM(high_count) AS total_high, SUM(med_count) AS total_med,
                   SUM(low_count) AS total_low
            FROM scans WHERE status = 'ok'
        """).fetchone()
        today = datetime.date.today().isoformat()
        today_row = conn.execute("""
            SELECT COUNT(*) AS scans_today, SUM(findings_count) AS pii_today
            FROM scans WHERE status = 'ok' AND timestamp LIKE ?
        """, (f"{today}%",)).fetchone()
        type_counts = {}
        for r in conn.execute(
            "SELECT types_found FROM scans WHERE status='ok' AND types_found IS NOT NULL"
        ).fetchall():
            try:
                for t in json.loads(r["types_found"]):
                    if t: type_counts[t] = type_counts.get(t, 0) + 1
            except Exception:
                pass
        file_types = conn.execute("""
            SELECT file_type, COUNT(*) as cnt FROM scans
            WHERE source='file' AND status='ok' AND file_type IS NOT NULL
            GROUP BY file_type ORDER BY cnt DESC
        """).fetchall()
        return {
            "total_scans":       row["total_scans"]        or 0,
            "total_pii_blocked": row["total_pii_blocked"]  or 0,
            "total_high":        row["total_high"]         or 0,
            "total_med":         row["total_med"]          or 0,
            "total_low":         row["total_low"]          or 0,
            "scans_today":       today_row["scans_today"]  or 0,
            "pii_today":         today_row["pii_today"]    or 0,
            "type_breakdown":    type_counts,
            "file_type_breakdown": {r["file_type"]: r["cnt"] for r in file_types},
        }
    finally:
        conn.close()

@app.get("/history")
def get_history(limit: int = 50):
    import json
    limit = min(limit, 200)
    conn = _get_db()
    try:
        rows = conn.execute(
            "SELECT * FROM scans ORDER BY timestamp DESC LIMIT ?", (limit,)
        ).fetchall()
        return {
            "records": [{
                "id":             r["id"],
                "timestamp":      r["timestamp"],
                "source":         r["source"],
                "file_name":      r["file_name"],
                "file_type":      r["file_type"],
                "findings_count": r["findings_count"],
                "types_found":    json.loads(r["types_found"] or "[]"),
                "high_count":     r["high_count"],
                "med_count":      r["med_count"],
                "low_count":      r["low_count"],
                "duration_ms":    r["duration_ms"],
                "status":         r["status"],
            } for r in rows],
            "total": len(rows),
        }
    finally:
        conn.close()

@app.delete("/history")
def clear_history():
    conn = _get_db()
    try:
        conn.execute("DELETE FROM scans")
        conn.commit()
        return {"status": "cleared"}
    finally:
        conn.close()