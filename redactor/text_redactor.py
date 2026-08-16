import re
import os
import shutil
import tempfile
import zipfile


# ─────────────────────────────────────────────────────────
# RUN MERGING — pure python-docx, no external script/process.
# Word splits text across multiple <w:r> runs (spell-check markers,
# autocorrect, etc). If a value straddles a run boundary the later
# XML string-replace won't find it as one contiguous string, so we
# coalesce adjacent runs that share identical formatting first.
# ─────────────────────────────────────────────────────────

def _runs_match_formatting(r1, r2):
    try:
        return (
            r1.font.bold == r2.font.bold
            and r1.font.italic == r2.font.italic
            and r1.font.underline == r2.font.underline
            and r1.font.size == r2.font.size
            and _run_color(r1) == _run_color(r2)
        )
    except Exception:
        return False


def _run_color(run):
    try:
        if run.font.color and run.font.color.rgb:
            return str(run.font.color.rgb)
    except Exception:
        pass
    return None


def _merge_paragraph_runs(paragraph):
    runs = paragraph.runs
    i = 0
    while i < len(runs) - 1:
        r1, r2 = runs[i], runs[i + 1]
        if _runs_match_formatting(r1, r2):
            r1.text = (r1.text or "") + (r2.text or "")
            r2._element.getparent().remove(r2._element)
            runs = paragraph.runs
        else:
            i += 1


def _merge_runs_docx(src_path, out_path):
    """Coalesce same-formatted adjacent runs across the whole document
    (body, tables, headers, footers) using only python-docx — no
    external binaries or scripts required."""
    import docx as docxlib

    doc = docxlib.Document(src_path)

    def _walk_paragraphs(paragraphs):
        for p in paragraphs:
            _merge_paragraph_runs(p)

    def _walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _walk_paragraphs(cell.paragraphs)
                    _walk_tables(cell.tables)

    _walk_paragraphs(doc.paragraphs)
    _walk_tables(doc.tables)

    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            try:
                _walk_paragraphs(part.paragraphs)
                _walk_tables(part.tables)
            except Exception:
                continue

    doc.save(out_path)


# ─────────────────────────────────────────────────────────
# TEXT REDACTOR  (plain text — used for typed messages)
# ─────────────────────────────────────────────────────────

def redact_text(original_text, findings):
    """
    Replace sensitive values in plain text with unique placeholders.
    Returns: (cleaned_text, summary, restore_map)
    """
    cleaned_text = original_text
    summary      = []
    restore_map  = {}
    type_counts  = {}

    # Replace from end → start so character positions don't shift
    sorted_findings = sorted(findings, key=lambda x: x.get("start", 0), reverse=True)

    for finding in sorted_findings:
        value   = finding["value"]
        replace = finding["replace"]
        ftype   = finding["type"]

        type_counts[ftype] = type_counts.get(ftype, 0) + 1
        unique_placeholder = replace.replace("]", f"_{type_counts[ftype]}]")
        restore_map[unique_placeholder] = value

        cleaned_text = re.sub(
            re.escape(value),
            unique_placeholder,
            cleaned_text,
            flags=re.IGNORECASE
        )
        summary.append({
            "type":     ftype,
            "original": value,
            "replace":  unique_placeholder,
            "risk":     finding.get("risk", "LOW"),
        })

    return cleaned_text, summary, restore_map


# ─────────────────────────────────────────────────────────
# DOCX REDACTOR  — direct XML replace, preserves 100% formatting
# ─────────────────────────────────────────────────────────

def redact_docx(src_path, findings, out_path):
    """
    Redact a .docx by doing direct XML string replacement.

    Handles:
      • Split runs  — merge_runs.py coalesces them first
      • XML-encoded chars  — &amp; &lt; &gt; &quot; &apos;
      • Tables, headers, footers, text boxes
      • Any font, color, layout — untouched

    Returns out_path on success, raises on failure.
    """
    tmp_dir = tempfile.mkdtemp(prefix="pg_docx_")
    try:
        # ── Step 1: merge split runs so values are contiguous in XML ──
        merged_src = os.path.join(tmp_dir, "merged.docx")
        try:
            _merge_runs_docx(src_path, merged_src)
        except Exception:
            # Merging is a best-effort optimization — if it fails for any
            # reason, fall back to the original file and continue. The
            # XML replace below still works for values that aren't split
            # across runs.
            shutil.copy2(src_path, merged_src)

        # ── Step 2: unzip ──────────────────────────────────────────────
        unzip_dir = os.path.join(tmp_dir, "unpacked")
        with zipfile.ZipFile(merged_src, "r") as z:
            z.extractall(unzip_dir)
        # Remove symlinks (untrusted docx)
        for root, dirs, files in os.walk(unzip_dir):
            for name in files + dirs:
                p = os.path.join(root, name)
                if os.path.islink(p):
                    os.unlink(p)

        # ── Step 3: build replacement map (XML-escaped search values) ──
        type_counts = {}
        replace_map = []   # list of (xml_escaped_value, placeholder)

        for finding in findings:
            value   = finding["value"]
            replace = finding["replace"]
            ftype   = finding["type"]
            type_counts[ftype] = type_counts.get(ftype, 0) + 1
            placeholder = replace.replace("]", f"_{type_counts[ftype]}]")

            # XML-encode the search value so it matches what's in the XML
            xml_value = (
                value
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&apos;")
            )
            replace_map.append((xml_value, placeholder, value))

        # ── Step 4: apply to every XML part in the docx ───────────────
        # Covers: document body, headers, footers, footnotes, endnotes,
        #         comments, text boxes (word/document.xml, header*.xml,
        #         footer*.xml, footnotes.xml, endnotes.xml, comments.xml)
        xml_targets = []
        for root, dirs, files in os.walk(unzip_dir):
            for fname in files:
                if fname.endswith(".xml"):
                    xml_targets.append(os.path.join(root, fname))

        for xml_path in xml_targets:
            try:
                with open(xml_path, "r", encoding="utf-8", errors="ignore") as f:
                    xml = f.read()

                changed = False
                for xml_value, placeholder, raw_value in replace_map:
                    if xml_value in xml:
                        xml = xml.replace(xml_value, placeholder)
                        changed = True
                    # Also try raw value in case it wasn't XML-encoded
                    elif raw_value in xml:
                        xml = xml.replace(raw_value, placeholder)
                        changed = True

                if changed:
                    with open(xml_path, "w", encoding="utf-8") as f:
                        f.write(xml)
            except Exception:
                continue  # skip non-text XML files (images etc.)

        # ── Step 5: rezip into output docx ────────────────────────────
        # Pure Python zipfile — works on Windows/Mac/Linux with no
        # external `zip`/`bash` binary required (that was the bug: on
        # machines without bash+zip on PATH, this step used to throw,
        # _cleaned_file_path stayed None, and the app silently fell back
        # to saving plain .txt instead of the cleaned .docx).
        if os.path.exists(out_path):
            os.remove(out_path)
        with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root, dirs, files in os.walk(unzip_dir):
                for fname in files:
                    abs_path = os.path.join(root, fname)
                    rel_path = os.path.relpath(abs_path, unzip_dir)
                    # Use forward slashes inside the zip regardless of OS
                    zf.write(abs_path, rel_path.replace(os.sep, "/"))

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return out_path


# ─────────────────────────────────────────────────────────
# PDF REDACTOR  — renders text layer, works on any PDF
# ─────────────────────────────────────────────────────────

def redact_pdf(src_path, findings, out_path):
    """
    Redact a PDF using pymupdf (fitz).

    Handles:
      • Text split across character-level spans  — search_for() uses
        rendered text layer, not raw stream
      • Unicode, ligatures, RTL text
      • Multi-page documents
      • Scanned PDFs (no text layer)  — graceful warning, saves unchanged

    Install:  pip install pymupdf

    Returns out_path on success, raises on failure.
    """
    try:
        import fitz
    except ImportError:
        raise ImportError(
            "pymupdf is required for PDF redaction. "
            "Install it with:  pip install pymupdf"
        )

    doc = fitz.open(src_path)
    type_counts = {}
    has_text    = any(page.get_text().strip() for page in doc)

    if not has_text:
        # Scanned PDF — no text layer to redact
        doc.close()
        shutil.copy2(src_path, out_path)
        raise ValueError(
            "This PDF appears to be a scanned image with no text layer. "
            "PrivacyGate cannot redact it automatically."
        )

    for finding in findings:
        value   = finding["value"]
        replace = finding["replace"]
        ftype   = finding["type"]
        type_counts[ftype] = type_counts.get(ftype, 0) + 1
        placeholder = replace.replace("]", f"_{type_counts[ftype]}]")

        for page in doc:
            # search_for searches the RENDERED text layer —
            # handles split spans, ligatures, encoding differences
            hits = page.search_for(value, quads=False)
            for rect in hits:
                # 1. White rectangle over original text
                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                # 2. Placeholder text in dark navy, same baseline
                fontsize = max(6, round(rect.height * 0.80))
                page.insert_text(
                    (rect.x0, rect.y1 - 1),
                    placeholder,
                    fontsize=fontsize,
                    color=(0.1, 0.2, 0.4),   # dark navy — readable on white
                )

    doc.save(out_path, garbage=4, deflate=True)
    doc.close()
    return out_path


# ─────────────────────────────────────────────────────────
# RESTORE  — put real values back after AI responds
# ─────────────────────────────────────────────────────────

def restore_text(ai_response, restore_map):
    """Replace placeholders in AI response back with real values."""
    restored_text   = ai_response
    restore_summary = []
    for placeholder, real_value in restore_map.items():
        if placeholder in restored_text:
            restored_text = restored_text.replace(placeholder, real_value)
            restore_summary.append({
                "placeholder": placeholder,
                "real_value":  real_value,
            })
    return restored_text, restore_summary


def get_redaction_stats(summary):
    stats = {"HIGH": 0, "MED": 0, "LOW": 0, "total": 0}
    for item in summary:
        risk = item.get("risk", "LOW")
        stats[risk] = stats.get(risk, 0) + 1
        stats["total"] += 1
    return stats