"""PrivacyGate cleaned-output verification.

Call verify_cleaned_file(original_path, cleaned_path, original_findings)
after the redactor has written the cleaned output and before returning it from
/scan-file. This module is intentionally fail-closed: unsupported or
unverifiable formats return ``uncertain`` rather than ``verified``.
"""
from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path
from typing import Any

from scanner.regex_scanner import scan_text_regex

try:
    from redactor.image_redactor import extract_text_image
except Exception:
    extract_text_image = None


def sha256_file(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normal(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", str(value).lower())


def _docx_text(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            parts.extend(p.text for p in part.paragraphs)
    return "\n".join(parts)


def _pdf_text(path: str) -> str | None:
    import fitz
    doc = fitz.open(path)
    try:
        text = "\n".join(page.get_text() for page in doc).strip()
    finally:
        doc.close()
    # Empty text means this is probably a scanned PDF. Without a page-image
    # OCR verifier, it is unsafe to call the output verified.
    return text or None


def _extract_clean_text(path: str, suffix: str) -> tuple[str | None, str]:
    if suffix in {".txt", ".csv", ".json", ".py", ".js", ".ts", ".html", ".md", ".yaml", ".yml", ".xml", ".log", ".ini", ".cfg", ".sql"}:
        return Path(path).read_text(encoding="utf-8", errors="ignore"), "text"
    if suffix == ".docx":
        return _docx_text(path), "docx"
    if suffix == ".pdf":
        return _pdf_text(path), "pdf"
    if suffix in {".xlsx", ".xlsm"}:
        from redactor.xlsx_redactor import extract_text_xlsx
        return extract_text_xlsx(path), "xlsx"
    if suffix == ".pptx":
        from redactor.pptx_redactor import extract_text_pptx
        return extract_text_pptx(path), "pptx"
    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".svg"}:
        if extract_text_image is None:
            return None, "image"
        return extract_text_image(path), "image"
    return None, "unsupported"


def verify_cleaned_file(original_path: str, cleaned_path: str,
                        original_findings: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    original_findings = original_findings or []
    result: dict[str, Any] = {
        "state": "uncertain",
        "reason": "verification_not_run",
        "remaining_pii": [],
        "original_sha256": sha256_file(original_path),
        "cleaned_sha256": sha256_file(cleaned_path),
        "original_size": os.path.getsize(original_path),
        "cleaned_size": os.path.getsize(cleaned_path),
    }

    if not os.path.exists(cleaned_path) or result["cleaned_size"] == 0:
        result.update(state="failed", reason="cleaned_output_missing_or_empty")
        return result
    if original_findings and result["original_sha256"] == result["cleaned_sha256"]:
        result.update(state="failed", reason="cleaned_output_is_identical_to_original")
        return result

    try:
        text, method = _extract_clean_text(cleaned_path, Path(cleaned_path).suffix.lower())
    except Exception as exc:
        result.update(state="uncertain", reason=f"rescan_error:{type(exc).__name__}")
        return result
    result["verification_method"] = method
    if text is None:
        result.update(state="uncertain", reason="format_cannot_be_rescanned")
        return result

    residual: list[dict[str, Any]] = []
    cleaned_normal = _normal(text)
    for finding in original_findings:
        value = str(finding.get("value", ""))
        normalized = _normal(value)
        if normalized and normalized in cleaned_normal:
            residual.append({"type": finding.get("type", "PII"), "value": value, "reason": "original_value_remains"})

    # Also scan the cleaned extraction. Any high/medium result means the
    # output still contains detectable sensitive content.
    for finding in scan_text_regex(text):
        if finding.get("risk", "LOW") in {"HIGH", "MED"}:
            if not any(_normal(finding.get("value", "")) == _normal(item.get("value", "")) for item in residual):
                residual.append({"type": finding.get("type", "PII"), "value": finding.get("value", ""), "reason": "rescan_detected"})

    result["remaining_pii"] = residual
    if residual:
        result.update(state="failed", reason="rescan_found_remaining_pii")
    else:
        result.update(state="verified", reason="rescan_found_no_high_or_medium_pii")
    return result
