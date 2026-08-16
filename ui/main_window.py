import tkinter as tk
import math, threading, random, webbrowser, os, shutil, tempfile, subprocess, urllib.parse, glob, re
from tkinter import filedialog
import customtkinter as ctk

try:
    import pyperclip
    HAS_PYPERCLIP = True
except ImportError:
    HAS_PYPERCLIP = False

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx
except ImportError:
    docx = None
try:
    import pytesseract
    from PIL import Image, ImageTk
except ImportError:
    pytesseract = None
    Image = None
    ImageTk = None

try:
    from PIL import Image as PilImage, ImageTk as PilImageTk
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from docx_native_render import render_docx_native
    HAS_NATIVE_DOCX_RENDER = True
except ImportError:
    HAS_NATIVE_DOCX_RENDER = False

try:
    from scanner.regex_scanner import scan_text_regex
    from scanner.ner_scanner import scan_text_ner
    from scanner.context_rules import apply_context_rules
    from redactor.text_redactor import redact_text, redact_docx, redact_pdf, restore_text
    from sender.autopaste import send_to_ai
    from storage.audit_log import write_log
    from storage.settings_store import load_settings
except ImportError:
    def scan_text_regex(t, document_type="General"): return []
    def scan_text_ner(t, document_type="General"): return []
    def apply_context_rules(f, t): return f
    def redact_text(t, f): return t, [], {}
    def redact_docx(s, f, o): return o
    def redact_pdf(s, f, o): return o
    def restore_text(t, m): return t, []
    def send_to_ai(tool, t): pass
    def write_log(d): pass
    def load_settings(): return {}

BG        = "#000000"
SIDEBAR_BG= "#050505"
ACCENT    = "#00f2ff"
ACCENT_DIM= "#006666"
TEXT_MAIN = "#ffffff"
TEXT_DIM  = "#a1a1aa"
BORDER    = "#27272a"
INPUT_BG  = "#18181b"
CARD      = "#0d0d0d"
GREEN     = "#00ff9d"
WHITE     = "#ffffff"
RED       = "#ff4444"
ORANGE    = "#ff9900"

PLACEHOLDER = "Message PrivacyGate..."

FONT_INPUT    = ("Segoe UI", 13)
FONT_BUBBLE   = ("Segoe UI", 12)
FONT_RESULT   = ("Segoe UI", 12)
FONT_LABEL    = ("Segoe UI", 10, "bold")
FONT_ARTIFACT = ("Segoe UI", 13)
FONT_BTN      = ("Segoe UI", 10, "bold")
FONT_HEADING  = ("Segoe UI", 11, "bold")
FONT_NAV      = ("Segoe UI", 13)
FONT_WELCOME  = ("Segoe UI", 28, "bold")
FONT_LOGO     = ("Segoe UI", 58, "bold")


def open_browser(url):
    chrome_paths = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        os.path.expanduser(r"~\AppData\Local\Google\Chrome\Application\chrome.exe"),
        "/usr/bin/google-chrome",
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    ]
    for p in chrome_paths:
        if os.path.exists(p):
            try:
                subprocess.Popen([p, url])
                return
            except Exception:
                continue
    webbrowser.open(url)


def _find_libreoffice():
    """Find LibreOffice / soffice binary on Windows, Mac, Linux."""
    candidates = [
        "libreoffice", "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/soffice",
    ]
    for c in candidates:
        try:
            result = subprocess.run([c, "--version"],
                                    capture_output=True, timeout=5)
            if result.returncode == 0:
                return c
        except Exception:
            continue
    return None


def _find_pdftoppm():
    """Find pdftoppm (Poppler). On Windows users may have it in PATH."""
    for c in ["pdftoppm"]:
        try:
            result = subprocess.run([c, "-v"], capture_output=True, timeout=3)
            if result.returncode == 0 or b"pdftoppm" in result.stderr:
                return c
        except Exception:
            continue
    return None


def doc_to_page_images(file_path, ftype, dpi=150):
    """
    Convert a docx or pdf to a list of PIL Images (one per page).

    Pipeline:
      docx  →  LibreOffice  →  PDF  →  pdftoppm OR pymupdf  →  [PIL Images]
      pdf   →                  PDF  →  pdftoppm OR pymupdf  →  [PIL Images]

    Returns list of PIL Image objects, or [] on failure.
    """
    if not HAS_PIL:
        return []

    tmp = tempfile.mkdtemp(prefix="pg_render_")
    images = []

    try:
        pdf_path = None

        # ── Step 1: get a PDF ──────────────────────────────────────────
        if ftype in ("docx", "doc"):
            lo = _find_libreoffice()
            if not lo:
                return []
            result = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", file_path, "--outdir", tmp],
                capture_output=True, timeout=60
            )
            # LibreOffice names the output after the input file
            base = os.path.splitext(os.path.basename(file_path))[0]
            pdf_path = os.path.join(tmp, base + ".pdf")
            if not os.path.exists(pdf_path):
                # Sometimes it picks the first .pdf in tmp
                pdfs = glob.glob(os.path.join(tmp, "*.pdf"))
                pdf_path = pdfs[0] if pdfs else None

        elif ftype == "pdf":
            pdf_path = file_path

        if not pdf_path or not os.path.exists(pdf_path):
            return []

        # ── Step 2: PDF → images ───────────────────────────────────────
        # Try pymupdf first (no external binary needed)
        try:
            import fitz
            doc = fitz.open(pdf_path)
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            for page in doc:
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = PilImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
            doc.close()
            return images
        except ImportError:
            pass

        # Fallback: pdftoppm
        pp = _find_pdftoppm()
        if pp:
            out_prefix = os.path.join(tmp, "page")
            subprocess.run(
                [pp, "-jpeg", "-r", str(dpi), pdf_path, out_prefix],
                capture_output=True, timeout=60
            )
            page_files = sorted(glob.glob(os.path.join(tmp, "page*.jpg")))
            for pf in page_files:
                images.append(PilImage.open(pf).copy())
            return images

        return []

    except Exception as e:
        print(f"doc_to_page_images error: {e}")
        return []
    finally:
        # Clean up temp — but only if we didn't use it as pdf_path directly
        try:
            shutil.rmtree(tmp, ignore_errors=True)
        except Exception:
            pass


class MainWindow(ctk.CTkFrame):
    def __init__(self, parent):
        super().__init__(parent, fg_color=BG)
        self.settings            = load_settings()
        self.findings            = []
        self.finding_vars        = []
        self.restore_map         = {}
        self.cleaned_text        = ""
        self.original_text       = ""
        self.cubes               = []
        self.particles           = []
        self.sensitivity         = tk.StringVar(value="Standard")
        self.attached_file_path  = None
        self.attached_file_type  = None
        self.attached_file_name  = None
        # Separate from the above: the "active" file currently being
        # scanned/cleaned. attached_file_* represents what's pending in
        # the compose box and gets wiped by _reset_input() right after
        # a file is read — but scanning/redaction/the artifact panel
        # need the file identity to survive well past that point, so
        # they read from these instead.
        self._active_file_path   = None
        self._active_file_type   = None
        self._active_file_name   = None
        self._placeholder_active = True
        self._scanning           = False
        self._artifact_open      = False
        self._last_file_name     = None
        self._cleaned_file_path  = None
        self._page_images        = []   # keep PIL ImageTk refs alive
        self.doc_type            = tk.StringVar(value="General")

        self._build_sidebar()

        self.main_container = tk.Frame(self, bg=BG)
        self.main_container.place(x=200, y=0, relwidth=1, relheight=1, width=-200)

        self.canvas = tk.Canvas(self.main_container, bg=BG, highlightthickness=0, bd=0)
        self.canvas.place(relx=0, rely=0, relwidth=1, relheight=1)

        self.output_inner  = tk.Frame(self.canvas, bg=BG)
        self.output_window = self.canvas.create_window(
            (0, 0), window=self.output_inner, anchor="nw"
        )

        self.artifact_panel = tk.Frame(
            self.main_container, bg=SIDEBAR_BG,
            highlightbackground=BORDER, highlightthickness=1
        )

        self._build_input_box()

        self.canvas.bind("<Configure>",       self._on_canvas_resize)
        self.output_inner.bind("<Configure>", self._on_content_resize)
        self.canvas.bind(
            "<MouseWheel>",
            lambda e: self.canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        )

        self.after(100, self._init_3d)
        self.after(500, self._show_welcome)

    # ─────────────────────────────────────────────────────────
    # 3D BACKGROUND
    # ─────────────────────────────────────────────────────────
    def _init_3d(self):
        W, H = self.canvas.winfo_width(), self.canvas.winfo_height()
        if W < 10:
            self.after(100, self._init_3d)
            return
        for _ in range(20):
            self.cubes.append({
                "x": random.uniform(50, W-50), "y": random.uniform(-H, H),
                "z": random.uniform(50, 400),
                "rx": random.random()*6, "ry": random.random()*6, "rz": random.random()*6,
                "dx": random.uniform(0.01,0.02), "dy": random.uniform(0.01,0.02),
                "dz": random.uniform(0.01,0.02),
                "speed": random.uniform(1.0,1.8), "size": random.uniform(20,40),
                "color": ACCENT if random.random()<0.3
                         else random.choice([WHITE,"#222222","#444444"]),
            })
        self.particles = [
            {"x": random.uniform(0,W), "y": random.uniform(0,H),
             "vx": random.uniform(-0.3,0.3), "vy": random.uniform(-0.3,0.3)}
            for _ in range(40)
        ]
        self._animate_bg()

    def _animate_bg(self):
        c = self.canvas
        W, H = c.winfo_width(), c.winfo_height()
        c.delete("3d")
        for p in self.particles:
            p["x"] = (p["x"]+p["vx"])%W; p["y"] = (p["y"]+p["vy"])%H
            c.create_oval(p["x"]-1,p["y"]-1,p["x"]+1,p["y"]+1,
                          fill="#111111",outline="",tags="3d")
        edges = [(0,1),(1,2),(2,3),(3,0),(4,5),(5,6),(6,7),(7,4),
                 (0,4),(1,5),(2,6),(3,7)]
        for cb in self.cubes:
            cb["rx"]+=cb["dx"]; cb["ry"]+=cb["dy"]; cb["rz"]+=cb["dz"]
            cb["y"]+=cb["speed"]
            if cb["y"]>H+80: cb["y"]=-80; cb["x"]=random.uniform(50,W-50)
            nodes=[]
            for i in range(8):
                x=cb["size"]*(1 if i&1 else -1)
                y=cb["size"]*(1 if i&2 else -1)
                z=cb["size"]*(1 if i&4 else -1)
                y,z=(y*math.cos(cb["rx"])-z*math.sin(cb["rx"]),
                     y*math.sin(cb["rx"])+z*math.cos(cb["rx"]))
                x,z=(x*math.cos(cb["ry"])+z*math.sin(cb["ry"]),
                    -x*math.sin(cb["ry"])+z*math.cos(cb["ry"]))
                x,y=(x*math.cos(cb["rz"])-y*math.sin(cb["rz"]),
                     x*math.sin(cb["rz"])+y*math.cos(cb["rz"]))
                f=400/(z+cb["z"])
                nodes.append((cb["x"]+x*f, cb["y"]+y*f))
            for e in edges:
                c.create_line(nodes[e[0]],nodes[e[1]],
                              fill=cb["color"],width=1.0,tags="3d")
        c.tag_lower("3d")
        self.after(40, self._animate_bg)

    # ─────────────────────────────────────────────────────────
    # WELCOME
    # ─────────────────────────────────────────────────────────
    def _show_welcome(self):
        for w in self.output_inner.winfo_children():
            w.destroy()
        self.canvas.delete("welcome")
        msg = random.choice([
            "What do you want to send to the AI today?",
            "PrivacyGate scans before anything leaves your device.",
            "Paste your message below — we clean it first.",
        ])
        W, H = self.canvas.winfo_width(), self.canvas.winfo_height()
        self.welcome_txt = self.canvas.create_text(
            W/2, H/2-50, text="", font=FONT_WELCOME,
            fill=TEXT_MAIN, width=600, justify="center", tags="welcome"
        )
        def type_in(i=0):
            if i<=len(msg):
                self.canvas.itemconfig(self.welcome_txt, text=msg[:i])
                self.after(45, lambda: type_in(i+1))
            else:
                self.after(2000, self._fade_to_logo)
        type_in()

    def _fade_to_logo(self):
        self.canvas.delete("welcome")
        W, H = self.canvas.winfo_width(), self.canvas.winfo_height()
        self.canvas.create_text(W/2, H/2-60, text="PrivacyGate",
                                font=FONT_LOGO, fill=ACCENT, tags="welcome")
        self.canvas.create_text(W/2, H/2+20, text="YOUR PRIVATE AI SAFETY LAYER",
                                font=("Segoe UI",12,"bold"), fill=TEXT_DIM, tags="welcome")

    # ─────────────────────────────────────────────────────────
    # SIDEBAR
    # ─────────────────────────────────────────────────────────
    def _build_sidebar(self):
        self.sidebar = tk.Frame(self, bg=SIDEBAR_BG, width=200)
        self.sidebar.place(x=0, y=0, width=200, relheight=1)
        tk.Label(self.sidebar, text="P", font=("Segoe UI",28,"bold"),
                 fg=ACCENT, bg=SIDEBAR_BG).pack(pady=(40,5))
        self.nav_btns = []
        for icon, txt, cmd in [
            ("⬡","Scan",     self._nav_scan),
            ("◈","Audit",    self._nav_audit),
            ("◎","Settings", self._nav_settings),
        ]:
            btn = ctk.CTkButton(
                self.sidebar, text=f"  {icon}   {txt}", font=FONT_NAV,
                fg_color="transparent", text_color=TEXT_DIM,
                hover_color="#111111", anchor="w", height=45,
                corner_radius=8, command=cmd
            )
            btn.pack(fill="x", padx=12, pady=4)
            self.nav_btns.append(btn)

        tk.Frame(self.sidebar, bg=BORDER, height=1).pack(fill="x", padx=16, pady=(20,10))
        ctk.CTkButton(
            self.sidebar, text="  ✦   New Chat", font=FONT_NAV,
            fg_color="#0d1f1f", text_color=ACCENT, hover_color="#0a1515",
            anchor="w", height=45, corner_radius=8,
            border_width=1, border_color=ACCENT_DIM,
            command=self._new_chat
        ).pack(fill="x", padx=12, pady=(0,4))

        tk.Label(self.sidebar, text="SENSITIVITY",
                 font=("Segoe UI",9,"bold"), fg="#444444",
                 bg=SIDEBAR_BG).pack(side="bottom", pady=(0,5))
        s_frame = tk.Frame(self.sidebar, bg=SIDEBAR_BG)
        s_frame.pack(side="bottom", fill="x", padx=15, pady=(0,20))
        for s in ["Standard","Deep"]:
            ctk.CTkRadioButton(
                s_frame, text=s, variable=self.sensitivity, value=s,
                font=("Segoe UI",11), fg_color=ACCENT,
                text_color=TEXT_DIM, border_color=BORDER, width=80
            ).pack(side="left")
        self._set_nav(0)

    # ─────────────────────────────────────────────────────────
    # INPUT BOX
    # ─────────────────────────────────────────────────────────
    def _build_input_box(self):
        self.input_container = ctk.CTkFrame(
            self.main_container, fg_color=INPUT_BG,
            corner_radius=28, border_width=1, border_color=BORDER
        )
        self.input_container.place(relx=0.5, rely=0.88, relwidth=0.75, anchor="center")

        self.file_chip_frame = tk.Frame(self.input_container, bg=INPUT_BG)

        pill_row = tk.Frame(self.input_container, bg=INPUT_BG)
        pill_row.pack(fill="x", padx=20, pady=(10,0))
        tk.Label(pill_row, text="Type:", font=("Segoe UI",9),
                 fg=TEXT_DIM, bg=INPUT_BG).pack(side="left", padx=(0,8))

        self._pill_btns = {}
        for dtype, icon in [("General","📝"),("Resume/CV","👤"),("Code","💻"),
                             ("Medical","🏥"),("Legal","⚖️"),("Financial","💰")]:
            btn = tk.Button(
                pill_row, text=f"{icon} {dtype}",
                font=("Segoe UI",9), relief="flat", padx=8, pady=3, cursor="hand2",
                fg=TEXT_DIM, bg="#1a1a1a",
                activebackground=ACCENT_DIM, activeforeground=WHITE,
                command=lambda d=dtype: self._select_doc_type(d)
            )
            btn.pack(side="left", padx=3)
            self._pill_btns[dtype] = btn
        self._select_doc_type("General")

        inner = tk.Frame(self.input_container, bg=INPUT_BG)
        inner.pack(fill="x", padx=20, pady=12)

        tk.Button(inner, text="+", font=("Segoe UI",18), fg=TEXT_DIM,
                  bg=INPUT_BG, relief="flat", cursor="hand2",
                  command=self._show_attach_menu).pack(side="left", padx=(0,15))

        self.message_input = tk.Text(
            inner, font=FONT_INPUT, fg=TEXT_DIM, bg=INPUT_BG,
            insertbackground=WHITE, relief="flat", height=1,
            highlightthickness=0, selectbackground=ACCENT, selectforeground=BG
        )
        self.message_input.pack(side="left", fill="x", expand=True)
        self.message_input.insert("1.0", PLACEHOLDER)

        self.message_input.bind("<FocusIn>",    self._clear_ph)
        self.message_input.bind("<FocusOut>",   self._restore_ph)
        self.message_input.bind("<KeyRelease>", self._auto_resize)
        self.message_input.bind("<<Paste>>",    self._on_paste)

        self.send_btn = ctk.CTkButton(
            inner, text="↑", font=("Segoe UI",18,"bold"),
            fg_color=ACCENT, hover_color=ACCENT_DIM,
            width=38, height=38, corner_radius=19,
            text_color=BG, command=self._on_send_click
        )
        self.send_btn.pack(side="right", padx=(10,0))

    # ─────────────────────────────────────────────────────────
    # PLACEHOLDER / INPUT HELPERS
    # ─────────────────────────────────────────────────────────
    def _clear_ph(self, e=None):
        if self._placeholder_active:
            self.message_input.delete("1.0","end")
            self.message_input.config(fg=WHITE)
            self.message_input.configure(height=1)
            self._placeholder_active = False

    def _restore_ph(self, e=None):
        if not self.message_input.get("1.0","end-1c").strip():
            self.message_input.delete("1.0","end")
            self.message_input.insert("1.0", PLACEHOLDER)
            self.message_input.config(fg=TEXT_DIM)
            self.message_input.configure(height=1)
            self._placeholder_active = True

    def _reset_input(self):
        self.message_input.delete("1.0","end")
        self.message_input.insert("1.0", PLACEHOLDER)
        self.message_input.config(fg=TEXT_DIM)
        self.message_input.configure(height=1)
        self._placeholder_active = True
        self.file_chip_frame.pack_forget()
        for w in self.file_chip_frame.winfo_children():
            w.destroy()
        self.attached_file_path = None
        self.attached_file_name = None
        self.attached_file_type = None

    def _get_input_text(self):
        if self._placeholder_active:
            return ""
        return self.message_input.get("1.0","end-1c").strip()

    def _auto_resize(self, e=None):
        if self._placeholder_active:
            return
        lines = int(self.message_input.index("end-1c").split(".")[0])
        self.message_input.configure(height=max(1, min(lines, 5)))

    def _select_doc_type(self, dtype):
        self.doc_type.set(dtype)
        for name, btn in self._pill_btns.items():
            btn.config(fg=BG if name==dtype else TEXT_DIM,
                       bg=ACCENT if name==dtype else "#1a1a1a")

    # ─────────────────────────────────────────────────────────
    # SEND
    # ─────────────────────────────────────────────────────────
    def _on_send_click(self):
        if self._scanning:
            return
        text = self._get_input_text()
        if not text and not self.attached_file_path:
            return
        if self._artifact_open:
            self._close_artifact()
        self.canvas.delete("welcome")
        self.original_text = text
        display = text if text else f"[Attached: {self.attached_file_name}]"
        self._add_user_bubble(display)
        self._reset_input()
        self._start_scan(from_file=False, text_override=text)

    # ─────────────────────────────────────────────────────────
    # PASTE
    # ─────────────────────────────────────────────────────────
    def _on_paste(self, e):
        self._clear_ph()
        try:
            result = subprocess.run(
                ["powershell","-command",
                 "Add-Type -Assembly System.Windows.Forms; "
                 "[System.Windows.Forms.Clipboard]::GetFileDropList()"],
                capture_output=True, text=True, timeout=2
            )
            files = [f.strip() for f in result.stdout.splitlines()
                     if f.strip() and os.path.exists(f.strip())]
            if files:
                self._process_file(files[0])
                return "break"
        except Exception:
            pass
        try:
            from PIL import ImageGrab
            img = ImageGrab.grabclipboard()
            if img:
                tmp = os.path.join(tempfile.gettempdir(),
                                   f"pasted_img_{random.randint(100,999)}.png")
                img.save(tmp)
                self._process_file(tmp, "Pasted Image.png")
                return "break"
        except Exception:
            pass
        return None

    # ─────────────────────────────────────────────────────────
    # FILE CHIP
    # ─────────────────────────────────────────────────────────
    def _show_file_chip(self, name, ftype):
        for w in self.file_chip_frame.winfo_children():
            w.destroy()
        self.file_chip_frame.pack(fill="x", padx=20, pady=(8,0))
        chip = tk.Frame(self.file_chip_frame, bg="#1a1a1a",
                        highlightbackground=ACCENT, highlightthickness=1)
        chip.pack(side="left", padx=5, pady=5)
        tk.Label(chip, text=f" {name[:35]} ", font=("Segoe UI",10),
                 fg=WHITE, bg="#1a1a1a").pack(side="left")
        tk.Button(chip, text="×", font=("Segoe UI",10), fg=TEXT_DIM,
                  bg="#1a1a1a", relief="flat",
                  command=self._clear_attachment).pack(side="left")

    def _clear_attachment(self):
        self.attached_file_path = None
        self.attached_file_name = None
        self.attached_file_type = None
        self.file_chip_frame.pack_forget()

    # ─────────────────────────────────────────────────────────
    # FILE PROCESSING
    # ─────────────────────────────────────────────────────────
    def _process_file(self, path, name=None):
        if not path:
            return
        ext  = path.lower().rsplit(".",1)[-1]
        name = name or os.path.basename(path)
        self.attached_file_path = path
        self.attached_file_type = "image" if ext in ("png","jpg","jpeg","bmp") else ext
        self.attached_file_name = name
        self._last_file_name    = name
        # Mirror into the "active" file slot so it survives _reset_input()
        self._active_file_path  = path
        self._active_file_type  = self.attached_file_type
        self._active_file_name  = name
        self._show_file_chip(name, self.attached_file_type)
        self.canvas.delete("welcome")
        status_bub = self._add_bubble(f"📄  Processing: {name}...", color=ACCENT, bg=CARD)
        threading.Thread(
            target=self._extract_content,
            args=(path, name, self.attached_file_type, status_bub),
            daemon=True
        ).start()

    def _extract_content(self, path, name, ftype, status_bub):
        text = ""
        try:
            if ftype == "pdf":
                if pdfplumber:
                    with pdfplumber.open(path) as pdf:
                        text = "\n".join(p.extract_text() or "" for p in pdf.pages)
                else:
                    text = "[Install pdfplumber to read PDF]"
            elif ftype in ("docx","doc"):
                if docx:
                    d = docx.Document(path)
                    parts = [p.text for p in d.paragraphs]
                    for table in d.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                parts.append(cell.text)
                    text = "\n".join(parts)
                else:
                    text = "[Install python-docx to read Word]"
            elif ftype == "image":
                if pytesseract:
                    text = pytesseract.image_to_string(Image.open(path))
                else:
                    text = "[Install pytesseract/Pillow for OCR]"
            else:
                with open(path,"r",encoding="utf-8",errors="ignore") as f:
                    text = f.read()
        except Exception as exc:
            text = f"Error reading file: {exc}"

        self.original_text = text.strip()

        def _after():
            try:
                status_bub.master.destroy()
            except Exception:
                pass
            self._add_bubble(f"✓  Read {name} successfully.", color=GREEN, bg=CARD)
            self._reset_input()
            self._start_scan(from_file=True)

        self.after(0, _after)

    # ─────────────────────────────────────────────────────────
    # ARTIFACT PANEL  ← THE BIG FIX
    # For docx/pdf: renders the actual cleaned file as page images
    # For plain text: styled scrollable text widget
    # ─────────────────────────────────────────────────────────
    def _show_artifact_panel(self, cleaned_text):
        self._artifact_open = True
        self.canvas.place(relx=0, rely=0, relwidth=0.55, relheight=1)
        self.input_container.place(relx=0.275, rely=0.88, relwidth=0.5, anchor="center")
        self.artifact_panel.place(relx=0.55, rely=0, relwidth=0.45, relheight=1)
        for w in self.artifact_panel.winfo_children():
            w.destroy()
        self._page_images = []  # clear old image refs

        panel_title = self._last_file_name or "Cleaned Message"

        # ── Header bar ────────────────────────────────────────
        header = tk.Frame(self.artifact_panel, bg="#111111")
        header.pack(fill="x")
        tk.Label(header, text=f"  {panel_title}", font=FONT_HEADING,
                 fg=ACCENT, bg="#111111", pady=10).pack(side="left")

        btn_f = tk.Frame(header, bg="#111111")
        btn_f.pack(side="right", padx=10)

        def download():
            if self._cleaned_file_path and os.path.exists(self._cleaned_file_path):
                ext = self._cleaned_file_path.rsplit(".",1)[-1]
                p = filedialog.asksaveasfilename(
                    defaultextension=f".{ext}",
                    initialfile=f"PrivacyGate_Cleaned.{ext}",
                    filetypes=[(f"{ext.upper()} files", f"*.{ext}"),("All files","*.*")]
                )
                if p:
                    shutil.copy2(self._cleaned_file_path, p)
                    self._add_bubble("✓  File saved.", color=GREEN, bg=CARD)
            else:
                p = filedialog.asksaveasfilename(
                    defaultextension=".txt",
                    initialfile="PrivacyGate_Cleaned.txt"
                )
                if p:
                    with open(p,"w",encoding="utf-8") as f:
                        f.write(cleaned_text)
                    self._add_bubble("✓  File saved.", color=GREEN, bg=CARD)

        tk.Button(btn_f, text=" ↓ Save ", font=FONT_BTN, fg=WHITE, bg="#1a3a1a",
                  relief="flat", padx=8, command=download).pack(side="left", padx=2, pady=6)
        tk.Button(btn_f, text=" Copy ", font=FONT_BTN, fg=WHITE, bg="#333333",
                  relief="flat", padx=8,
                  command=lambda: (self.clipboard_clear(),
                                   self.clipboard_append(cleaned_text),
                                   self.update())).pack(side="left", padx=2, pady=6)
        tk.Button(btn_f, text=" Share ▾ ", font=FONT_BTN, fg=BG, bg=ACCENT,
                  relief="flat", padx=8,
                  command=lambda: self._show_share_menu(cleaned_text)).pack(side="left", padx=2, pady=6)
        tk.Button(btn_f, text=" × ", font=("Segoe UI",12), fg=TEXT_DIM, bg="#111111",
                  relief="flat", command=self._close_artifact).pack(side="left", padx=5, pady=6)

        # ── Content area ──────────────────────────────────────
        content = tk.Frame(self.artifact_panel, bg="#e8e8e8")
        content.pack(fill="both", expand=True)

        ftype = (self._active_file_type or "").lower()

        if self._cleaned_file_path and ftype in ("docx", "doc") and HAS_NATIVE_DOCX_RENDER:
            # FAST PATH — native colored render straight from docx XML.
            # No LibreOffice/Poppler required, so this works everywhere.
            try:
                render_docx_native(content, self._cleaned_file_path)
            except Exception as e:
                print(f"native docx render failed, falling back: {e}")
                for w in content.winfo_children():
                    w.destroy()
                self._render_via_image_pipeline(content, ftype, cleaned_text)

        elif self._cleaned_file_path and ftype in ("docx", "doc", "pdf"):
            self._render_via_image_pipeline(content, ftype, cleaned_text)

        elif self._cleaned_file_path and ftype == "image" and HAS_PIL:
            self._show_image_preview(content, self._cleaned_file_path)

        else:
            # Plain text fallback (typed messages, images, code)
            self._show_text_fallback(content, cleaned_text)

    def _show_image_preview(self, content, path):
        """Show an attached/pasted image inline in the artifact panel."""
        try:
            img = PilImage.open(path)
            img.thumbnail((500, 800))
            tk_img = PilImageTk.PhotoImage(img)
            self._page_images = [tk_img]  # keep reference alive
            tk.Label(content, image=tk_img, bg="#e8e8e8").pack(expand=True, pady=20)
        except Exception as e:
            print(f"image preview failed: {e}")
            self._show_text_fallback(content, "[Could not preview image]")

    def _render_via_image_pipeline(self, content, ftype, cleaned_text):
        """LibreOffice -> PDF -> image fallback (used for PDFs, or if native
        docx rendering isn't available)."""
        spinner = tk.Label(content, text="⏳  Rendering document...",
                           font=("Segoe UI",12), fg="#555555", bg="#e8e8e8")
        spinner.pack(expand=True)

        def _render_thread():
            imgs = doc_to_page_images(self._cleaned_file_path, ftype, dpi=150)
            self.after(0, lambda: self._display_page_images(content, spinner, imgs, cleaned_text))

        threading.Thread(target=_render_thread, daemon=True).start()

        self.after(10, lambda: self.canvas.yview_moveto(1.0))

    def _display_page_images(self, content, spinner, imgs, cleaned_text):
        """Called on main thread after rendering finishes."""
        try:
            spinner.destroy()
        except Exception:
            pass

        if not imgs:
            # Render failed — fall back to text
            self._show_text_fallback(content, cleaned_text)
            return

        # Scrollable canvas showing page images
        outer = tk.Frame(content, bg="#888888")
        outer.pack(fill="both", expand=True)

        cvs = tk.Canvas(outer, bg="#888888", highlightthickness=0)
        sb  = tk.Scrollbar(outer, orient="vertical", command=cvs.yview)
        cvs.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        cvs.pack(side="left", fill="both", expand=True)

        inner = tk.Frame(cvs, bg="#888888")
        win   = cvs.create_window((0,0), window=inner, anchor="nw")

        def _on_resize(e):
            cvs.configure(scrollregion=cvs.bbox("all"))
            cvs.itemconfig(win, width=e.width)
        inner.bind("<Configure>", _on_resize)
        cvs.bind("<Configure>",   lambda e: cvs.itemconfig(win, width=e.width))
        cvs.bind("<MouseWheel>",
                 lambda e: cvs.yview_scroll(int(-1*(e.delta/120)),"units"))

        # Get panel width for scaling
        self.artifact_panel.update_idletasks()
        panel_w = self.artifact_panel.winfo_width() - 30  # scrollbar + padding
        if panel_w < 100:
            panel_w = 550

        for img in imgs:
            # Scale image to fit panel width
            img_w, img_h = img.size
            scale = panel_w / img_w
            new_w = int(img_w * scale)
            new_h = int(img_h * scale)
            img_resized = img.resize((new_w, new_h), PilImage.LANCZOS)
            tk_img = PilImageTk.PhotoImage(img_resized)
            self._page_images.append(tk_img)   # keep reference!

            lbl = tk.Label(inner, image=tk_img, bg="#888888", cursor="arrow")
            lbl.pack(pady=6)

    def _show_text_fallback(self, content, cleaned_text):
        """Plain scrollable text widget — used for typed messages / code."""
        txt = tk.Text(
            content, font=FONT_ARTIFACT, fg=TEXT_MAIN, bg=SIDEBAR_BG,
            relief="flat", padx=24, pady=20, wrap="word",
            highlightthickness=0
        )
        txt.insert("1.0", cleaned_text)
        txt.config(state="disabled")
        sb = tk.Scrollbar(content, command=txt.yview, bg=SIDEBAR_BG)
        txt.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        txt.pack(fill="both", expand=True)

    def _close_artifact(self):
        self._artifact_open = False
        self._page_images   = []
        self.canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.input_container.place(relx=0.5, rely=0.88, relwidth=0.75, anchor="center")
        self.artifact_panel.place_forget()

    # ─────────────────────────────────────────────────────────
    # SHARE MENU
    # ─────────────────────────────────────────────────────────
    def _show_share_menu(self, text_to_share):
        share = tk.Toplevel(self)
        share.overrideredirect(True)
        share.configure(bg="#0a0a0a")
        share.attributes("-topmost", True)
        ap = self.artifact_panel
        x  = ap.winfo_rootx() + 20
        y  = ap.winfo_rooty() + 50
        share.geometry(f"280x300+{x}+{y}")

        hdr = tk.Frame(share, bg="#111111")
        hdr.pack(fill="x")
        tk.Label(hdr, text="  Share to AI Tool", font=FONT_HEADING,
                 fg=ACCENT, bg="#111111", pady=12).pack(side="left")
        tk.Button(hdr, text="×", font=("Segoe UI",12), fg=TEXT_DIM,
                  bg="#111111", relief="flat",
                  command=share.destroy).pack(side="right", padx=10)

        tk.Label(share, text="  Text is copied + pre-filled in browser",
                 font=("Segoe UI",9), fg="#555555", bg="#0a0a0a",
                 anchor="w").pack(fill="x", padx=10, pady=(6,2))

        encoded = urllib.parse.quote(text_to_share, safe="")
        tools = [
            ("Claude",  "🤖", "https://claude.ai/new"),
            ("ChatGPT", "💬", f"https://chatgpt.com/?q={encoded}"),
            ("Gemini",  "✦",  "https://gemini.google.com/app?hl=en"),
            ("Copilot", "🪟", f"https://copilot.microsoft.com/?q={encoded}"),
        ]

        def _do_share(url):
            if HAS_PYPERCLIP:
                try: pyperclip.copy(text_to_share)
                except Exception: pass
            self.clipboard_clear()
            self.clipboard_append(text_to_share)
            self.update()
            share.destroy()
            self.after(150, lambda: open_browser(url))

        for lbl, icon, url in tools:
            row = tk.Frame(share, bg="#0a0a0a")
            row.pack(fill="x", padx=10, pady=2)
            tk.Button(row, text=f" {icon}  {lbl}", font=("Segoe UI",12),
                      fg=WHITE, bg="#111111", activebackground=ACCENT_DIM,
                      activeforeground=WHITE, relief="flat", anchor="w",
                      pady=9, cursor="hand2",
                      command=lambda u=url: _do_share(u)).pack(fill="x")

        tk.Frame(share, bg=BORDER, height=1).pack(fill="x", padx=16, pady=8)
        tk.Button(share, text="  📋  Copy only (no browser)",
                  font=("Segoe UI",10), fg=TEXT_DIM, bg="#0a0a0a",
                  relief="flat", anchor="w", pady=6,
                  command=lambda: (
                      self.clipboard_clear(),
                      self.clipboard_append(text_to_share),
                      self.update(), share.destroy(),
                      self._add_bubble("✓  Cleaned text copied to clipboard.",
                                       color=GREEN, bg=CARD)
                  )).pack(fill="x", padx=10)

    # ─────────────────────────────────────────────────────────
    # SCAN FLOW
    # ─────────────────────────────────────────────────────────
    def _start_scan(self, from_file=False, text_override=None):
        if from_file:
            self._add_user_bubble(f"[Attached: {self._last_file_name or 'file'}]")
        self._scanning = True
        self._add_divider()
        scan_bub = self._add_bubble("🔍  Scanning locally...", color=TEXT_DIM, bg=CARD)
        self.after(20, lambda: self.canvas.yview_moveto(1.0))
        threading.Thread(
            target=self._run_scan,
            args=(self.original_text, scan_bub),
            daemon=True
        ).start()

    # ─────────────────────────────────────────────────────────
    # SUPPLEMENTAL PATTERN SCAN — LinkedIn / GitHub profile links.
    # The external scanner.regex_scanner module doesn't catch these, so
    # we add them here directly rather than relying on a file we can't
    # see or edit.
    # ─────────────────────────────────────────────────────────
    _LINKEDIN_RE = re.compile(
        r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_%]+/?',
        re.IGNORECASE
    )
    _GITHUB_RE = re.compile(
        r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9\-_]+/?',
        re.IGNORECASE
    )

    def _scan_additional_patterns(self, text, existing_values):
        """Find LinkedIn/GitHub profile URLs not already caught by the
        main scanners. existing_values = lowercased set of values already
        found, so we don't add duplicates."""
        extra = []
        for regex, ftype, label in (
            (self._LINKEDIN_RE, "LINKEDIN", "[LINKEDIN]"),
            (self._GITHUB_RE,   "GITHUB",   "[GITHUB]"),
        ):
            for m in regex.finditer(text):
                value = m.group(0)
                if value.lower() in existing_values:
                    continue
                existing_values.add(value.lower())
                extra.append({
                    "value":   value,
                    "type":    ftype,
                    "replace": label,
                    "risk":    "MED",
                    "start":   m.start(),
                })
        return extra

    def _run_scan(self, text, scan_bub):
        dtype    = self.doc_type.get()
        findings = []
        findings.extend(scan_text_regex(text, document_type=dtype))
        findings.extend(scan_text_ner(text,   document_type=dtype))
        if self.sensitivity.get() == "Deep":
            findings = apply_context_rules(findings, text)

        existing_values = {f.get("value","").lower() for f in findings}
        findings.extend(self._scan_additional_patterns(text, existing_values))

        if dtype == "Resume/CV":
            RESUME_SENSITIVE = {
                "EMAIL","PHONE","AADHAAR","PAN","PASSPORT",
                "BANK_ACCOUNT","SSN","CREDIT_CARD","NI_NUMBER",
                "TAX_ID","PASSWORD","API_KEY","SECRET",
                "LINKEDIN","GITHUB",
            }
            findings = [f for f in findings
                        if f.get("type","").upper() in RESUME_SENSITIVE]
        elif dtype == "Code":
            CODE_SENSITIVE = {
                "API_KEY","SECRET","PASSWORD","TOKEN",
                "EMAIL","IP_ADDRESS","DATABASE_URL","PRIVATE_KEY",
                "LINKEDIN","GITHUB",
            }
            findings = [f for f in findings
                        if f.get("type","").upper() in CODE_SENSITIVE]
        elif dtype == "Medical":
            MEDICAL_SKIP = {"ORG","PERSON","DATE","CARDINAL"}
            findings = [f for f in findings
                        if f.get("ner_label","") not in MEDICAL_SKIP]

        self.findings = findings
        self.after(0, lambda: self._show_findings(scan_bub))

    def _show_findings(self, scan_bub):
        self._scanning = False
        try:
            scan_bub.master.destroy()
        except Exception:
            pass

        if not self.findings:
            self._add_bubble("✓  No sensitive data found.", color=GREEN, bg=CARD)
            self.cleaned_text = self.original_text
            self._cleaned_file_path = self._active_file_path
            self._show_artifact_panel(self.cleaned_text)
            return

        score = max(0, 100 - len(self.findings)*10)
        self._add_bubble(
            f"Privacy Score: {score}%",
            color=ORANGE if score<75 else GREEN, bg=CARD
        )
        self.finding_vars = []
        for f in self.findings:
            self._add_finding_row(f)
        self._show_finding_actions()
        self.after(20, lambda: self.canvas.yview_moveto(1.0))

    def _add_finding_row(self, finding):
        var = tk.BooleanVar(value=True)
        self.finding_vars.append((var, finding))
        row = tk.Frame(self.output_inner, bg=INPUT_BG,
                       highlightbackground=ACCENT, highlightthickness=1)
        row.pack(fill="x", padx=48, pady=3)

        def _toggle(v=var, r=row, b=None):
            v.set(not v.get())
            checked = v.get()
            r.config(highlightbackground=ACCENT if checked else BORDER)
            if b:
                b.config(fg=BG if checked else TEXT_DIM,
                         bg=ACCENT if checked else "#333333")

        btn = tk.Button(
            row, text="✓", font=("Segoe UI",9,"bold"),
            fg=BG, bg=ACCENT, relief="flat", padx=6
        )
        btn.config(command=lambda v=var, r=row, b=btn: _toggle(v, r, b))
        btn.pack(side="left", padx=10, pady=8)

        tk.Label(row, text=finding.get("type",""), font=FONT_LABEL,
                 fg=TEXT_DIM, bg=INPUT_BG, width=14, anchor="w").pack(side="left")
        tk.Label(row, text=finding.get("value","")[:40], font=FONT_RESULT,
                 fg=WHITE, bg=INPUT_BG).pack(side="left", fill="x", expand=True)
        tk.Label(row, text=f"→ {finding.get('replace','')}", font=FONT_RESULT,
                 fg=ACCENT, bg=INPUT_BG).pack(side="right", padx=12)

    def _show_finding_actions(self):
        f = tk.Frame(self.output_inner, bg=BG)
        f.pack(fill="x", padx=48, pady=10)
        tk.Button(
            f, text=" Clean & Continue → ", font=("Segoe UI",12,"bold"),
            fg=BG, bg=ACCENT, relief="flat", padx=14, pady=7,
            command=self._clean_and_continue
        ).pack(side="left")

    def _clean_and_continue(self):
        checked = [f for v, f in self.finding_vars if v.get()]
        self.cleaned_text, _, self.restore_map = redact_text(self.original_text, checked)
        self._add_divider()
        self._add_bubble("✓  Cleaned.", color=GREEN, bg=CARD)

        src   = self._active_file_path
        ftype = (self._active_file_type or "").lower()
        out_path = src  # default: pass the original through (e.g. images,
                         # or docx/pdf when nothing was checked for redaction)

        if src and checked:
            tmp_dir = tempfile.gettempdir()
            if ftype in ("docx","doc"):
                out_path = os.path.join(tmp_dir, "privacygate_cleaned.docx")
                try:
                    redact_docx(src, checked, out_path)
                except Exception as e:
                    self._add_bubble(f"⚠  docx redaction error: {e}", color=ORANGE, bg=CARD)
                    out_path = None
            elif ftype == "pdf":
                out_path = os.path.join(tmp_dir, "privacygate_cleaned.pdf")
                try:
                    redact_pdf(src, checked, out_path)
                except Exception as e:
                    self._add_bubble(f"⚠  PDF redaction error: {e}", color=ORANGE, bg=CARD)
                    out_path = None
            # image / other types: out_path stays as src — text-layer PII
            # (from OCR) is already stripped out of cleaned_text above;
            # we don't attempt pixel-level redaction on the image itself.

        self._cleaned_file_path = out_path
        self._show_artifact_panel(self.cleaned_text)

    # ─────────────────────────────────────────────────────────
    # CHAT BUBBLES / DIVIDER
    # ─────────────────────────────────────────────────────────
    def _add_bubble(self, text, color=WHITE, bg=CARD):
        f = tk.Frame(self.output_inner, bg=BG)
        f.pack(fill="x", padx=32, pady=3)
        bub = tk.Frame(f, bg=bg, padx=14, pady=10)
        bub.pack(fill="x")
        tk.Label(bub, text=text, font=FONT_BUBBLE, fg=color, bg=bg,
                 anchor="w", justify="left", wraplength=500).pack(fill="x")
        return bub

    def _add_user_bubble(self, text):
        f = tk.Frame(self.output_inner, bg=BG)
        f.pack(fill="x", padx=32, pady=3)
        bub = tk.Frame(f, bg="#0d1b1b", padx=14, pady=10)
        bub.pack(side="right")
        tk.Label(bub, text=text[:300], font=FONT_BUBBLE,
                 fg=ACCENT, bg="#0d1b1b", wraplength=460).pack()

    def _add_divider(self):
        tk.Frame(self.output_inner, bg=BORDER, height=1).pack(fill="x", padx=32, pady=8)

    # ─────────────────────────────────────────────────────────
    # CANVAS / SCROLL
    # ─────────────────────────────────────────────────────────
    def _on_canvas_resize(self, e):
        self.canvas.itemconfig(self.output_window, width=e.width)

    def _on_content_resize(self, e):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    # ─────────────────────────────────────────────────────────
    # NAV
    # ─────────────────────────────────────────────────────────
    def _nav_scan(self):
        self._set_nav(0); self._close_artifact(); self._show_welcome()

    def _nav_audit(self):
        self._set_nav(1); self._close_artifact(); self.canvas.delete("welcome")

    def _nav_settings(self):
        self._set_nav(2); self._close_artifact(); self.canvas.delete("welcome")

    def _set_nav(self, idx):
        for i, b in enumerate(self.nav_btns):
            b.configure(
                text_color=TEXT_MAIN if i==idx else TEXT_DIM,
                fg_color="#0d0d0d"   if i==idx else "transparent"
            )

    def _new_chat(self):
        if self._artifact_open:
            self._close_artifact()
        for w in self.output_inner.winfo_children():
            w.destroy()
        self.findings         = []
        self.finding_vars     = []
        self.restore_map      = {}
        self.cleaned_text     = ""
        self.original_text    = ""
        self._scanning        = False
        self._last_file_name  = None
        self._cleaned_file_path = None
        self._active_file_path = None
        self._active_file_type = None
        self._active_file_name = None
        self._page_images     = []
        self._reset_input()
        self._select_doc_type("General")
        self._set_nav(0)
        self._show_welcome()

    # ─────────────────────────────────────────────────────────
    # ATTACH MENU
    # ─────────────────────────────────────────────────────────
    def _show_attach_menu(self):
        m = tk.Menu(self, tearoff=0, bg=CARD, fg=WHITE,
                    font=("Segoe UI",11),
                    activebackground=ACCENT, activeforeground=BG, bd=0)
        m.add_command(label="📄  Document", command=self._pick_and_process)
        m.add_command(label="🖼️  Image",   command=self._pick_and_process)
        x = self.winfo_rootx() + 220
        y = self.winfo_rooty() + self.winfo_height() - 150
        m.tk_popup(x, y)

    def _pick_and_process(self):
        path = filedialog.askopenfilename()
        if path:
            self._process_file(path)